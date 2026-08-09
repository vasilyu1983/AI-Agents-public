#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import sys
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path
from typing import Any

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
NS = {"w": W_NS, "r": R_NS}

EXTRA_FIELDS = {"headers", "footers", "hyperlinks", "comments", "images"}


def _require_python_docx():
    try:
        from docx import Document  # type: ignore

        return Document
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency: python-docx. Install with: pip install python-docx"
        ) from exc


def _normalize_text(value: str) -> str:
    return " ".join(value.split())


def _read_xml(zip_file: zipfile.ZipFile, name: str) -> ET.Element | None:
    try:
        with zip_file.open(name) as file:
            return ET.fromstring(file.read())
    except KeyError:
        return None


def _element_text(element: ET.Element) -> str:
    return _normalize_text("".join(element.itertext()))


def _collect_paragraphs(root: ET.Element) -> list[str]:
    paragraphs: list[str] = []
    for paragraph in root.findall(".//w:p", NS):
        text = _element_text(paragraph)
        if text:
            paragraphs.append(text)
    return paragraphs


def _build_relationship_map(zip_file: zipfile.ZipFile, rel_name: str) -> dict[str, str]:
    root = _read_xml(zip_file, rel_name)
    if root is None:
        return {}

    mapping: dict[str, str] = {}
    for rel in root.findall(f"{{{PKG_REL_NS}}}Relationship"):
        rel_id = rel.attrib.get("Id")
        target = rel.attrib.get("Target")
        if rel_id and target:
            mapping[rel_id] = target
    return mapping


def _extract_hyperlinks(zip_file: zipfile.ZipFile) -> list[dict[str, Any]]:
    document_root = _read_xml(zip_file, "word/document.xml")
    if document_root is None:
        return []

    rels = _build_relationship_map(zip_file, "word/_rels/document.xml.rels")
    links: list[dict[str, Any]] = []
    for hyperlink in document_root.findall(".//w:hyperlink", NS):
        rel_id = hyperlink.attrib.get(f"{{{R_NS}}}id")
        anchor = hyperlink.attrib.get(f"{{{W_NS}}}anchor")
        text = _element_text(hyperlink)
        target = rels.get(rel_id)
        if text or target or anchor:
            links.append(
                {
                    "text": text or None,
                    "target": target,
                    "anchor": anchor,
                    "relationship_id": rel_id,
                }
            )
    return links


def _extract_comments(zip_file: zipfile.ZipFile) -> list[dict[str, Any]]:
    comments_root = _read_xml(zip_file, "word/comments.xml")
    if comments_root is None:
        return []

    comments: list[dict[str, Any]] = []
    for comment in comments_root.findall("w:comment", NS):
        comments.append(
            {
                "id": comment.attrib.get(f"{{{W_NS}}}id"),
                "author": comment.attrib.get(f"{{{W_NS}}}author"),
                "initials": comment.attrib.get(f"{{{W_NS}}}initials"),
                "date": comment.attrib.get(f"{{{W_NS}}}date"),
                "text": _element_text(comment) or None,
            }
        )
    return comments


def _extract_part_collection(zip_file: zipfile.ZipFile, prefix: str) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for name in sorted(part for part in zip_file.namelist() if part.startswith(prefix) and part.endswith(".xml")):
        root = _read_xml(zip_file, name)
        if root is None:
            continue
        items.append({"part": name, "paragraphs": _collect_paragraphs(root)})
    return items


def _extract_images(zip_file: zipfile.ZipFile) -> list[dict[str, Any]]:
    images: list[dict[str, Any]] = []
    for info in sorted(zip_file.infolist(), key=lambda item: item.filename):
        if not info.filename.startswith("word/media/"):
            continue
        images.append(
            {
                "path": info.filename,
                "filename": Path(info.filename).name,
                "bytes": info.file_size,
                "extension": Path(info.filename).suffix.lower(),
            }
        )
    return images


def extract_docx(docx_path: Path, include: set[str] | None = None) -> dict[str, Any]:
    include = include or set()
    Document = _require_python_docx()
    doc = Document(str(docx_path))

    paragraphs = [p.text for p in doc.paragraphs]
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        tables.append([[cell.text for cell in row.cells] for row in table.rows])

    props = doc.core_properties
    core_properties = {
        "title": props.title,
        "subject": props.subject,
        "author": props.author,
        "category": props.category,
        "comments": props.comments,
        "created": props.created.isoformat() if props.created else None,
        "modified": props.modified.isoformat() if props.modified else None,
    }

    payload: dict[str, Any] = {
        "path": str(docx_path),
        "core_properties": core_properties,
        "paragraphs": paragraphs,
        "tables": tables,
    }

    if not include:
        return payload

    with zipfile.ZipFile(docx_path) as zip_file:
        if "headers" in include:
            payload["headers"] = _extract_part_collection(zip_file, "word/header")
        if "footers" in include:
            payload["footers"] = _extract_part_collection(zip_file, "word/footer")
        if "hyperlinks" in include:
            payload["hyperlinks"] = _extract_hyperlinks(zip_file)
        if "comments" in include:
            payload["comments"] = _extract_comments(zip_file)
        if "images" in include:
            payload["images"] = _extract_images(zip_file)

    return payload


def _parse_include_args(raw_values: list[str]) -> set[str]:
    requested = {value.strip().lower() for value in raw_values if value.strip()}
    if "all" in requested:
        return set(EXTRA_FIELDS)
    invalid = requested - EXTRA_FIELDS
    if invalid:
        raise ValueError(
            f"Invalid --include value(s): {', '.join(sorted(invalid))}. "
            f"Expected one or more of: {', '.join(sorted(EXTRA_FIELDS))}, all"
        )
    return requested


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description="Extract text and tables from a .docx into JSON.")
    parser.add_argument("docx", type=Path, help="Path to a .docx file")
    parser.add_argument("--out", type=Path, help="Output JSON path (defaults to stdout)")
    parser.add_argument(
        "--include",
        nargs="*",
        default=[],
        help="Optional extras: headers footers hyperlinks comments images or all",
    )
    args = parser.parse_args(argv)

    if not args.docx.exists():
        print(f"File not found: {args.docx}", file=sys.stderr)
        return 2

    try:
        include = _parse_include_args(args.include)
        payload = extract_docx(args.docx, include=include)
    except Exception as exc:
        print(str(exc), file=sys.stderr)
        return 2

    output = json.dumps(payload, indent=2, ensure_ascii=False)
    if args.out:
        args.out.write_text(output, encoding="utf-8")
    else:
        print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
